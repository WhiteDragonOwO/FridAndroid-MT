#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取docx文件内容到txt文件
"""

from docx import Document
import sys


def extract_docx_content(docx_path, output_path):
    """
    提取docx文件内容到txt文件
    
    Args:
        docx_path: docx文件路径
        output_path: 输出txt文件路径
    """
    try:
        # 打开docx文件
        doc = Document(docx_path)
        
        # 提取所有段落内容
        content = []
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text.strip())
        
        # 提取所有表格内容
        for table in doc.tables:
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_content.append(cell.text.strip())
                if row_content:
                    content.append('\t'.join(row_content))
        
        # 写入到txt文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
        
        print(f"成功提取docx内容到 {output_path}")
        return True
    except Exception as e:
        print(f"提取docx内容失败: {e}")
        return False


if __name__ == "__main__":
    # 检查命令行参数
    if len(sys.argv) != 3:
        print("用法: python extract_docx.py <docx_file_path> <output_txt_file_path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2]
    
    extract_docx_content(docx_path, output_path)